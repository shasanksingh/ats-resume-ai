from pathlib import Path
from typing import Dict, List

from app.core.config import UPLOAD_DIR
from app.services.llm.resume_optimizer import optimize_resume
from app.services.parser.resume_parser import extract_resume_data


OUTPUT_DIR = UPLOAD_DIR / "optimized"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def rebuild_resume(
    resume_text: str,
    jd_text: str,
    rag_context: str,
    basename: str,
    output_format: str = "docx",
    structured_resume: Dict | None = None,
    confirmed_keywords: List[str] | None = None,
) -> Dict:
    structured = structured_resume or extract_resume_data(resume_text)
    result = optimize_resume(
        resume_text,
        jd_text,
        rag_context,
        include_local_model=False,
        structured_resume=structured,
        confirmed_keywords=confirmed_keywords,
    )
    structured = result["structured_resume"]
    optimized = result["optimization"]
    output_format = output_format.lower()

    if output_format == "pdf":
        path = OUTPUT_DIR / f"{Path(basename).stem}_optimized.pdf"
        build_pdf(path, structured, optimized)
    else:
        path = OUTPUT_DIR / f"{Path(basename).stem}_optimized.docx"
        build_docx(path, structured, optimized)

    return {
        "filename": path.name,
        "path": f"optimized/{path.name}",
        "download_url": f"/files/optimized/{path.name}",
        "format": output_format,
    }


def build_docx(path: Path, structured: Dict, optimized: Dict) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.42)
    section.bottom_margin = Inches(0.42)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9)
    normal.paragraph_format.space_after = Pt(1.5)

    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = heading.add_run(structured.get("name") or "Optimized Resume")
    name_run.bold = True
    name_run.font.size = Pt(20)

    contact = contact_line(structured)
    if contact:
        paragraph = document.add_paragraph(contact)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(5)

    add_docx_section(
        document,
        "Professional Summary",
        [optimized.get("improved_summary") or structured.get("summary", "")],
    )
    add_docx_skills(document, structured)
    add_docx_experience(document, structured, optimized)
    add_docx_projects(document, structured, optimized)
    add_docx_education(document, structured)
    add_docx_section(document, "Certifications", structured.get("certifications", []))
    document.save(path)


def add_docx_heading(document, title: str) -> None:
    from docx.shared import Pt

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(10.5)
    paragraph.add_run("\n" + ("_" * 105))


def add_docx_section(document, title: str, items: List[str]) -> None:
    cleaned = [str(item).strip() for item in items if str(item).strip()]
    if not cleaned:
        return
    add_docx_heading(document, title)
    for item in cleaned:
        document.add_paragraph(item)


def add_docx_skills(document, structured: Dict) -> None:
    groups = structured.get("skill_groups") or {}
    if not groups and structured.get("skills"):
        groups = {"Core Skills": structured["skills"]}
    if not groups:
        return
    add_docx_heading(document, "Skills")
    for label, values in groups.items():
        paragraph = document.add_paragraph()
        label_run = paragraph.add_run(f"{label}: ")
        label_run.bold = True
        paragraph.add_run(", ".join(values))


def add_docx_experience(document, structured: Dict, optimized: Dict) -> None:
    experience = structured.get("experience", [])
    if not experience:
        return
    add_docx_heading(document, "Experience")
    rewrites = rewrite_lookup(optimized)
    for role in experience:
        add_docx_entry_header(
            document,
            role.get("company") or role.get("title") or "",
            role.get("dates") or "",
            role.get("title") or "",
        )
        for bullet in role.get("bullets", []):
            document.add_paragraph(rewrites.get(bullet, bullet), style="List Bullet")


def add_docx_projects(document, structured: Dict, optimized: Dict) -> None:
    projects = structured.get("projects", [])
    if not projects:
        return
    add_docx_heading(document, "Projects")
    rewrites = rewrite_lookup(optimized)
    for project in projects:
        add_docx_entry_header(
            document,
            project.get("name") or "",
            project.get("date") or "",
            project.get("technologies") or "",
        )
        for bullet in project.get("bullets", []):
            document.add_paragraph(rewrites.get(bullet, bullet), style="List Bullet")


def add_docx_education(document, structured: Dict) -> None:
    education = structured.get("education_entries", [])
    if not education:
        add_docx_section(document, "Education", structured.get("education", []))
        return
    add_docx_heading(document, "Education")
    for item in education:
        add_docx_entry_header(
            document,
            item.get("institution") or "",
            item.get("dates") or "",
            item.get("degree") or "",
            item.get("details") or "",
        )


def add_docx_entry_header(
    document,
    left: str,
    right: str,
    subtitle: str = "",
    detail: str = "",
) -> None:
    from docx.shared import Inches, Pt

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.8))
    left_run = paragraph.add_run(left)
    left_run.bold = True
    paragraph.add_run(f"\t{right}")
    if subtitle or detail:
        meta = document.add_paragraph()
        meta.paragraph_format.space_after = Pt(1)
        meta.paragraph_format.tab_stops.add_tab_stop(Inches(6.8))
        subtitle_run = meta.add_run(subtitle)
        subtitle_run.italic = True
        meta.add_run(f"\t{detail}")


def build_pdf(path: Path, structured: Dict, optimized: Dict) -> None:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        HRFlowable,
        KeepTogether,
        ListFlowable,
        ListItem,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ResumeTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=18,
        leading=19,
        alignment=TA_CENTER,
        spaceAfter=1,
    )
    contact_style = ParagraphStyle(
        "Contact",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=7.8,
        leading=9,
        alignment=TA_CENTER,
        spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "ResumeBody",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=8.3,
        leading=9.8,
        spaceAfter=1.5,
    )
    meta_style = ParagraphStyle(
        "Meta",
        parent=body_style,
        fontName="Helvetica-Oblique",
        textColor=colors.HexColor("#333333"),
    )
    entry_style = ParagraphStyle(
        "Entry",
        parent=body_style,
        fontName="Helvetica-Bold",
        fontSize=8.7,
        leading=10,
    )
    section_style = ParagraphStyle(
        "Section",
        parent=body_style,
        fontName="Helvetica-Bold",
        fontSize=10.5,
        leading=11.4,
        spaceBefore=3,
        spaceAfter=1,
    )
    bullet_style = ParagraphStyle(
        "Bullet",
        parent=body_style,
        leftIndent=0,
        firstLineIndent=0,
        spaceAfter=0.2,
    )

    story = [
        Paragraph(html_escape(structured.get("name") or "Optimized Resume"), title_style),
    ]
    contact = contact_line(structured)
    if contact:
        story.append(Paragraph(html_escape(contact), contact_style))

    summary = optimized.get("improved_summary") or structured.get("summary", "")
    if summary:
        add_pdf_section(story, "Professional Summary", section_style)
        story.append(Paragraph(html_escape(summary), body_style))

    groups = structured.get("skill_groups") or {}
    if not groups and structured.get("skills"):
        groups = {"Core Skills": structured["skills"]}
    if groups:
        add_pdf_section(story, "Skills", section_style)
        for label, values in groups.items():
            story.append(
                Paragraph(
                    f"<b>{html_escape(label)}:</b> {html_escape(', '.join(values))}",
                    body_style,
                )
            )

    rewrites = rewrite_lookup(optimized)
    experience = structured.get("experience", [])
    if experience:
        add_pdf_section(story, "Experience", section_style)
        for role in experience:
            block = [
                pdf_entry_table(
                    role.get("company") or role.get("title") or "",
                    role.get("dates") or "",
                    role.get("title") or "",
                    "",
                    entry_style,
                    meta_style,
                    Table,
                    TableStyle,
                    colors,
                ),
                pdf_bullets(role.get("bullets", []), rewrites, bullet_style, ListFlowable, ListItem, Paragraph),
            ]
            story.append(KeepTogether([item for item in block if item]))

    projects = structured.get("projects", [])
    if projects:
        add_pdf_section(story, "Projects", section_style)
        for project in projects:
            block = [
                pdf_entry_table(
                    project.get("name") or "",
                    project.get("date") or "",
                    project.get("technologies") or "",
                    "",
                    entry_style,
                    meta_style,
                    Table,
                    TableStyle,
                    colors,
                ),
                pdf_bullets(project.get("bullets", []), rewrites, bullet_style, ListFlowable, ListItem, Paragraph),
            ]
            story.append(KeepTogether([item for item in block if item]))

    education = structured.get("education_entries", [])
    if education:
        add_pdf_section(story, "Education", section_style)
        for item in education:
            story.append(
                pdf_entry_table(
                    item.get("institution") or "",
                    item.get("dates") or "",
                    item.get("degree") or "",
                    item.get("details") or "",
                    entry_style,
                    meta_style,
                    Table,
                    TableStyle,
                    colors,
                )
            )

    certifications = structured.get("certifications", [])
    if certifications:
        add_pdf_section(story, "Certifications", section_style)
        story.append(pdf_bullets(certifications, {}, bullet_style, ListFlowable, ListItem, Paragraph))

    document = SimpleDocTemplate(
        str(path),
        pagesize=LETTER,
        leftMargin=0.38 * inch,
        rightMargin=0.38 * inch,
        topMargin=0.3 * inch,
        bottomMargin=0.3 * inch,
        title=structured.get("name") or "Optimized Resume",
    )
    document.build(story)


def add_pdf_section(story, title: str, section_style) -> None:
    from reportlab.lib import colors
    from reportlab.platypus import HRFlowable, Paragraph

    story.append(Paragraph(html_escape(title), section_style))
    story.append(HRFlowable(width="100%", thickness=0.45, color=colors.HexColor("#555555"), spaceAfter=1.5))


def pdf_entry_table(
    left: str,
    right: str,
    subtitle: str,
    detail: str,
    entry_style,
    meta_style,
    Table,
    TableStyle,
    colors,
):
    from reportlab.platypus import Paragraph

    rows = [[Paragraph(html_escape(left), entry_style), Paragraph(html_escape(right), entry_style)]]
    if subtitle or detail:
        rows.append([Paragraph(html_escape(subtitle), meta_style), Paragraph(html_escape(detail), meta_style)])
    table = Table(rows, colWidths=["76%", "24%"], hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ALIGN", (1, 0), (1, -1), "RIGHT"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0.5),
                ("TEXTCOLOR", (0, 0), (-1, -1), colors.black),
            ]
        )
    )
    return table


def pdf_bullets(items, rewrites, bullet_style, ListFlowable, ListItem, Paragraph):
    if not items:
        return None
    return ListFlowable(
        [
            ListItem(
                Paragraph(html_escape(rewrites.get(item, item)), bullet_style),
                leftIndent=8,
            )
            for item in items
        ],
        bulletType="bullet",
        start="-",
        leftIndent=13,
        bulletFontName="Helvetica",
        bulletFontSize=4.5,
        spaceBefore=0,
        spaceAfter=1,
    )


def contact_line(structured: Dict) -> str:
    links = structured.get("links") or {}
    values = [
        structured.get("phone"),
        structured.get("email"),
        "LinkedIn" if links.get("linkedin") else "",
        "GitHub" if links.get("github") else "",
        "Portfolio" if links.get("portfolio") else "",
    ]
    return " | ".join(str(value) for value in values if value)


def rewrite_lookup(optimized: Dict) -> Dict[str, str]:
    return {
        item.get("original", ""): item.get("improved", "")
        for item in optimized.get("improved_bullets", [])
        if item.get("original") and item.get("improved")
    }


def html_escape(value: object) -> str:
    from xml.sax.saxutils import escape

    return escape(str(value))
